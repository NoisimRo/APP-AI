"""Export service for generating DOCX, PDF, and MD files from training materials."""

import io
import os
import re
from datetime import datetime

from app.core.logging import get_logger

logger = get_logger(__name__)


def _strip_markdown(text: str) -> list[dict]:
    """Parse markdown text into structured elements for DOCX generation.

    Returns list of dicts with keys: type (heading2, heading3, paragraph, bullet, numbered), text, bold_ranges.
    """
    elements = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            elements.append({"type": "heading3", "text": stripped[4:]})
        elif stripped.startswith("## "):
            elements.append({"type": "heading2", "text": stripped[3:]})
        elif stripped.startswith("# "):
            elements.append({"type": "heading1", "text": stripped[2:]})
        elif stripped.startswith("- ") or stripped.startswith("* "):
            elements.append({"type": "bullet", "text": stripped[2:]})
        elif re.match(r"^\d+\.\s", stripped):
            text_content = re.sub(r"^\d+\.\s", "", stripped)
            elements.append({"type": "numbered", "text": text_content})
        elif stripped.startswith("|") and stripped.endswith("|"):
            elements.append({"type": "table_row", "text": stripped})
        else:
            elements.append({"type": "paragraph", "text": stripped})
    return elements


def _add_formatted_text(paragraph, text: str):
    """Add text with bold/italic markdown formatting to a docx paragraph."""
    # Split by bold markers
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif "*" in part:
            # Handle italic
            sub_parts = re.split(r"(\*.*?\*)", part)
            for sub in sub_parts:
                if sub.startswith("*") and sub.endswith("*") and not sub.startswith("**"):
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub)
        else:
            paragraph.add_run(part)


def export_markdown(content: str, title: str) -> bytes:
    """Export content as markdown bytes."""
    header = f"# {title}\n\n"
    header += f"*Generat de ExpertAP TrainingAP — {datetime.now().strftime('%d.%m.%Y %H:%M')}*\n\n---\n\n"
    full_content = header + content
    return full_content.encode("utf-8")


def export_docx(content: str, title: str, metadata: dict | None = None) -> bytes:
    """Export content as DOCX bytes using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata subtitle
    meta_parts = []
    if metadata:
        if metadata.get("tip_name"):
            meta_parts.append(f"Tip: {metadata['tip_name']}")
        if metadata.get("nivel"):
            nivel_labels = {"usor": "Ușor", "mediu": "Mediu", "dificil": "Dificil", "foarte_dificil": "Foarte Dificil"}
            meta_parts.append(f"Nivel: {nivel_labels.get(metadata['nivel'], metadata['nivel'])}")
    meta_parts.append(f"Generat: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    subtitle = doc.add_paragraph(" | ".join(meta_parts))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # spacer

    # Parse and add content
    elements = _strip_markdown(content)
    for elem in elements:
        if elem["type"] == "heading1":
            doc.add_heading(elem["text"], level=1)
        elif elem["type"] == "heading2":
            doc.add_heading(elem["text"], level=2)
        elif elem["type"] == "heading3":
            doc.add_heading(elem["text"], level=3)
        elif elem["type"] == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, elem["text"])
        elif elem["type"] == "numbered":
            p = doc.add_paragraph(style="List Number")
            _add_formatted_text(p, elem["text"])
        elif elem["type"] == "table_row":
            # Skip table formatting rows (separators)
            if re.match(r"^\|[\s\-:|]+\|$", elem["text"]):
                continue
            # Simple table rows rendered as paragraph
            cells = [c.strip() for c in elem["text"].split("|")[1:-1]]
            p = doc.add_paragraph()
            p.add_run(" | ".join(cells)).font.size = Pt(10)
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, elem["text"])

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph("—")
    footer_para.add_run("\nGenerat de ExpertAP TrainingAP | expertap.ro")
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(160, 160, 160)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _find_dejavu_font() -> str:
    """Find DejaVu Sans font path on the system."""
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path.rsplit("/", 1)[0]  # return directory
    raise FileNotFoundError("DejaVu Sans font not found. Install fonts-dejavu-core.")


def export_pdf(content: str, title: str, metadata: dict | None = None) -> bytes:
    """Export content as PDF bytes using fpdf2 with Unicode font support."""
    from fpdf import FPDF

    font_dir = _find_dejavu_font()

    class TrainingPDF(FPDF):
        def header(self):
            self.set_font("DejaVu", "I", 8)
            self.set_text_color(160, 160, 160)
            self.cell(0, 10, "ExpertAP TrainingAP", align="R", new_x="LMARGIN", new_y="NEXT")

        def footer(self):
            self.set_y(-15)
            self.set_font("DejaVu", "I", 8)
            self.set_text_color(160, 160, 160)
            self.cell(0, 10, f"Pagina {self.page_no()}/{{nb}}", align="C")

    pdf = TrainingPDF()
    pdf.add_font("DejaVu", "", os.path.join(font_dir, "DejaVuSans.ttf"))
    pdf.add_font("DejaVu", "B", os.path.join(font_dir, "DejaVuSans-Bold.ttf"))
    pdf.add_font("DejaVu", "I", os.path.join(font_dir, "DejaVuSansMono-Oblique.ttf"))
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    pdf.set_font("DejaVu", "B", 18)
    pdf.set_text_color(30, 41, 59)
    pdf.multi_cell(0, 10, title, align="C")
    pdf.ln(3)

    # Metadata
    meta_parts = []
    if metadata:
        if metadata.get("tip_name"):
            meta_parts.append(f"Tip: {metadata['tip_name']}")
        if metadata.get("nivel"):
            nivel_labels = {"usor": "Ușor", "mediu": "Mediu", "dificil": "Dificil", "foarte_dificil": "Foarte Dificil"}
            meta_parts.append(f"Nivel: {nivel_labels.get(metadata['nivel'], metadata['nivel'])}")
    meta_parts.append(f"Generat: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    pdf.set_font("DejaVu", "I", 9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 6, " | ".join(meta_parts), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)

    # Content
    elements = _strip_markdown(content)
    for elem in elements:
        if elem["type"] in ("heading1", "heading2"):
            pdf.ln(4)
            pdf.set_font("DejaVu", "B", 14)
            pdf.set_text_color(30, 41, 59)
            heading_text = elem["text"].replace("**", "")
            pdf.multi_cell(0, 7, heading_text)
            pdf.ln(2)
        elif elem["type"] == "heading3":
            pdf.ln(2)
            pdf.set_font("DejaVu", "B", 12)
            pdf.set_text_color(51, 65, 85)
            heading_text = elem["text"].replace("**", "")
            pdf.multi_cell(0, 6, heading_text)
            pdf.ln(1)
        elif elem["type"] == "bullet":
            pdf.set_font("DejaVu", "", 10)
            pdf.set_text_color(30, 41, 59)
            clean_text = elem["text"].replace("**", "")
            pdf.cell(8)
            pdf.multi_cell(0, 5, f"\u2022 {clean_text}")
            pdf.ln(1)
        elif elem["type"] == "numbered":
            pdf.set_font("DejaVu", "", 10)
            pdf.set_text_color(30, 41, 59)
            clean_text = elem["text"].replace("**", "")
            pdf.cell(8)
            pdf.multi_cell(0, 5, f"  {clean_text}")
            pdf.ln(1)
        elif elem["type"] == "table_row":
            if re.match(r"^\|[\s\-:|]+\|$", elem["text"]):
                continue
            pdf.set_font("DejaVu", "", 9)
            pdf.set_text_color(30, 41, 59)
            cells = [c.strip() for c in elem["text"].split("|")[1:-1]]
            clean_cells = [c.replace("**", "") for c in cells]
            pdf.multi_cell(0, 5, " | ".join(clean_cells))
        else:
            pdf.set_font("DejaVu", "", 10)
            pdf.set_text_color(30, 41, 59)
            clean_text = elem["text"].replace("**", "")
            pdf.multi_cell(0, 5, clean_text)
            pdf.ln(1)

    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer.getvalue()
